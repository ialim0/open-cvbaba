"""
Word Document Generation Service using python-docx.

Converts HTML documents to Word (.docx) files.
"""

import logging
import re
from io import BytesIO
from typing import Optional, List

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class WordService:
    """Service for converting HTML documents to Word."""
    
    @classmethod
    def html_to_word(
        cls,
        html_content: str,
        pages: Optional[List[int]] = None,
        page_indices: Optional[List[int]] = None
    ) -> bytes:
        """
        Convert HTML content to Word document bytes.
        
        Args:
            html_content: The HTML string to convert
            pages: List of 0-based page indices to keep. If None, keep all.
            page_indices: Alias for pages.
        """
        target_pages = pages if pages is not None else page_indices
        try:
            doc = Document()
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Find all pdf-page divs or process the whole body
            pdf_pages = soup.find_all('div', class_='pdf-page')
            
            if pdf_pages:
                # Filter pages if requested
                if target_pages is not None:
                    indexes_to_keep = set(target_pages)
                    pdf_pages = [p for i, p in enumerate(pdf_pages) if i in indexes_to_keep]
                
                if not pdf_pages:
                    logger.warning("No pages left after filtering")
                    # Should we error or return empty doc? Empty doc seems safer.
                
                for i, page in enumerate(pdf_pages):
                    cls._process_element(doc, page)
                    # Add page break between pages (but not after the last one)
                    if i < len(pdf_pages) - 1:
                        doc.add_page_break()
            else:
                # Process body content directly
                body = soup.find('body') or soup
                cls._process_element(doc, body)
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            word_bytes = buffer.getvalue()
            
            logger.info(f"Word document generated successfully ({len(word_bytes)} bytes)")
            return word_bytes
            
        except Exception as e:
            logger.error(f"Word generation failed: {e}", exc_info=True)
            raise
    
    @classmethod
    def create_word_document(
        cls,
        html_content: str,
        pages: Optional[List[int]] = None,
        page_indices: Optional[List[int]] = None
    ) -> BytesIO:
        """
        Create a Word document and return as a BytesIO stream.
        """
        word_bytes = cls.html_to_word(html_content, pages=pages, page_indices=page_indices)
        buffer = BytesIO(word_bytes)
        buffer.seek(0)
        return buffer
    
    @classmethod
    def _process_element(cls, doc: Document, element, paragraph=None):
        """Recursively process HTML elements and add them to the Word document."""
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text and paragraph:
                paragraph.add_run(text)
            return
        
        tag_name = element.name if element.name else ''
        
        # Skip style and script tags
        if tag_name in ('style', 'script', 'meta', 'link'):
            return
        
        # Handle different HTML elements
        if tag_name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag_name[1])
            text = element.get_text(strip=True)
            if text:
                heading = doc.add_heading(text, level=min(level, 9))
        
        elif tag_name == 'p':
            p = doc.add_paragraph()
            cls._process_inline_content(p, element)
        
        elif tag_name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                cls._process_inline_content(p, li)
        
        elif tag_name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                cls._process_inline_content(p, li)
        
        elif tag_name == 'br':
            if paragraph:
                paragraph.add_run('\n')
            else:
                doc.add_paragraph()
        
        elif tag_name == 'hr':
            # Add a horizontal line as a paragraph with a border
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif tag_name == 'table':
            cls._process_table(doc, element)
        
        elif tag_name == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            cls._process_inline_content(p, element)
        
        elif tag_name in ('div', 'section', 'article', 'main', 'header', 'footer', 'nav', 'aside'):
            # Container elements - process children
            for child in element.children:
                cls._process_element(doc, child, paragraph)
        
        elif tag_name == 'span':
            # Inline element - add to current paragraph or create new one
            if paragraph:
                cls._process_inline_content(paragraph, element)
            else:
                p = doc.add_paragraph()
                cls._process_inline_content(p, element)
        
        elif tag_name in ('strong', 'b', 'em', 'i', 'u', 'a'):
            # These should be handled in inline content processing
            if paragraph:
                cls._process_inline_content(paragraph, element)
        
        else:
            # For unknown elements, process children
            for child in element.children:
                cls._process_element(doc, child, paragraph)
    
    @classmethod
    def _process_inline_content(cls, paragraph, element):
        """Process inline content and add runs to the paragraph."""
        if isinstance(element, NavigableString):
            text = str(element)
            if text.strip():
                paragraph.add_run(text)
            return
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    paragraph.add_run(text)
            else:
                tag_name = child.name if child.name else ''
                text = child.get_text()
                
                if tag_name in ('strong', 'b'):
                    run = paragraph.add_run(text)
                    run.bold = True
                elif tag_name in ('em', 'i'):
                    run = paragraph.add_run(text)
                    run.italic = True
                elif tag_name == 'u':
                    run = paragraph.add_run(text)
                    run.underline = True
                elif tag_name == 'a':
                    run = paragraph.add_run(text)
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                elif tag_name == 'br':
                    paragraph.add_run('\n')
                elif tag_name == 'span':
                    cls._process_inline_content(paragraph, child)
                elif tag_name in ('strong', 'b', 'em', 'i', 'u'):
                    # Handle nested inline elements
                    cls._process_inline_content(paragraph, child)
                else:
                    # For other inline elements, just add the text
                    if text:
                        paragraph.add_run(text)
    
    @classmethod
    def _process_table(cls, doc: Document, table_element):
        """Process an HTML table and add it to the Word document."""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Count max columns
        max_cols = 0
        for row in rows:
            cells = row.find_all(['th', 'td'])
            max_cols = max(max_cols, len(cells))
        
        if max_cols == 0:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    word_cell = table.rows[row_idx].cells[col_idx]
                    text = cell.get_text(strip=True)
                    word_cell.text = text
                    
                    # Bold for header cells
                    if cell.name == 'th':
                        for paragraph in word_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    @classmethod
    def generate_filename(cls, title: str) -> str:
        """Generate a safe filename from the document title."""
        # Remove or replace unsafe characters
        safe_title = re.sub(r'[<>:"/\\|?*]', '', title)
        safe_title = safe_title.strip()
        
        # Limit length
        if len(safe_title) > 100:
            safe_title = safe_title[:100]
        
        # Default if empty
        if not safe_title:
            safe_title = "document"
        
        return f"{safe_title}.docx"
